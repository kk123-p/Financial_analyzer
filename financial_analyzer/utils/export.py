"""
数据导出模块 - 支持 Excel/CSV/JSON/PDF/Word 格式
"""
import os
import json
from datetime import datetime
from pathlib import Path

import pandas as pd

from ..config import PDF_FONT_PATHS, PDF_FONT_NAME, USER_DATA_DIR, AUTO_SAVE_DIR
from ..logging_config import get_logger

logger = get_logger(__name__)

# 数据类型 → 中文 sheet 名映射
SHEET_NAMES = {
    "daily": "日线行情",
    "daily_basic": "每日指标",
    "basic": "基本信息",
    "income": "利润表",
    "balance": "资产负债表",
    "cashflow": "现金流量表",
    "financial": "财务指标",
}


class DataExporter:
    """数据导出器"""

    @staticmethod
    def export_data(data: dict, file_path: str):
        """根据文件扩展名自动选择导出格式"""
        ext = Path(file_path).suffix.lower()
        if ext == ".xlsx":
            DataExporter.save_to_excel(data, file_path)
        elif ext == ".json":
            DataExporter.save_to_json(data, file_path)
        else:
            DataExporter.save_to_csv(data, file_path)

    @staticmethod
    def auto_save(data: dict, stock_code: str, save_dir: Path = None) -> str | None:
        """获取数据后自动保存到本地目录

        默认保存到 D:/FinancialAnalyzerData/{stock_code}/{YYYYMMDD}/
        格式: CSV（每种数据类型一个文件）

        Returns:
            保存目录路径，或 None
        """
        if save_dir is None:
            save_dir = AUTO_SAVE_DIR

        today = datetime.now().strftime("%Y%m%d")
        stock_dir = save_dir / stock_code / today
        stock_dir.mkdir(parents=True, exist_ok=True)

        saved = []
        for dtype, df in data.items():
            if df is not None and isinstance(df, pd.DataFrame) and not df.empty:
                file_path = stock_dir / f"{dtype}.csv"
                try:
                    df.to_csv(file_path, index=False, encoding="utf-8-sig")
                    saved.append(dtype)
                except Exception as e:
                    logger.error(f"自动保存 {dtype} 失败: {e}")

        if saved:
            logger.info(f"数据已自动保存到 {stock_dir} ({', '.join(saved)})")
            return str(stock_dir)
        return None

    @staticmethod
    def save_to_excel(data: dict, file_path: str, analysis_result: str = "",
                      stock_code: str = "", analysis_type: str = ""):
        """保存所有数据到一个 Excel 文件（每种数据类型一个 sheet）

        Args:
            data: 数据字典，key 为数据类型，value 为 DataFrame
            file_path: 保存路径
            analysis_result: 可选的分析结果文本
            stock_code: 股票代码（用于 sheet 头部）
            analysis_type: 分析类型
        """
        with pd.ExcelWriter(file_path, engine="openpyxl") as writer:
            for dtype, df in data.items():
                if df is not None and isinstance(df, pd.DataFrame) and not df.empty:
                    sheet_name = SHEET_NAMES.get(dtype, dtype)[:31]  # Excel sheet 名最长 31 字符
                    df.to_excel(writer, sheet_name=sheet_name, index=False)

            if analysis_result:
                lines = analysis_result.split("\n")
                pd.DataFrame({"分析内容": lines}).to_excel(
                    writer, sheet_name="分析报告", index=False
                )

        logger.info(f"Excel 文件已保存: {file_path}")

    @staticmethod
    def save_to_csv(data: dict, file_path: str):
        """保存为 CSV 文件（仅保存日线数据）"""
        if "daily" in data and data["daily"] is not None and not data["daily"].empty:
            data["daily"].to_csv(file_path, index=False, encoding="utf-8-sig")
        else:
            raise ValueError("没有可保存的数据")
        logger.info(f"CSV 文件已保存: {file_path}")

    @staticmethod
    def save_to_json(data: dict, file_path: str, analysis_result: str = ""):
        """保存为 JSON 文件"""
        save_data = {}
        for dtype, df in data.items():
            if df is not None and isinstance(df, pd.DataFrame) and not df.empty:
                save_data[dtype] = df.to_dict("records")
        if analysis_result:
            save_data["analysis"] = analysis_result
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2, default=str)
        logger.info(f"JSON 文件已保存: {file_path}")

    @staticmethod
    def save_analysis_to_txt(analysis_result: str, file_path: str,
                             stock_code: str = "", analysis_type: str = ""):
        """保存分析结果为文本文件"""
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("=" * 60 + "\n")
            f.write("                   财务分析报告\n")
            f.write("=" * 60 + "\n\n")
            f.write(f"股票代码: {stock_code}\n")
            f.write(f"分析类型: {analysis_type}\n")
            f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")
            f.write(analysis_result)
        logger.info(f"TXT 文件已保存: {file_path}")

    @staticmethod
    def save_analysis_to_pdf(analysis_result: str, file_path: str,
                             stock_code: str = "", analysis_type: str = ""):
        """保存分析结果为 PDF 文件（支持中文）"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            font_registered = False
            for font_path in PDF_FONT_PATHS:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, font_path))
                        font_registered = True
                        logger.info(f"已注册中文字体: {font_path}")
                        break
                    except Exception as e:
                        logger.warning(f"字体注册失败 {font_path}: {e}")

            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            font_name = PDF_FONT_NAME if font_registered else "Helvetica"
            c.setFont(font_name, 12)

            c.drawString(1 * inch, height - 1 * inch, "财务分析报告")
            c.setFont(font_name, 10)
            c.drawString(1 * inch, height - 1.2 * inch, f"股票代码: {stock_code}")
            c.drawString(1 * inch, height - 1.4 * inch, f"分析类型: {analysis_type}")
            c.drawString(1 * inch, height - 1.6 * inch,
                         f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            y = height - 2 * inch
            lines = analysis_result.split("\n")
            for line in lines:
                if y < 1 * inch:
                    c.showPage()
                    c.setFont(font_name, 10)
                    y = height - 1 * inch
                c.drawString(1 * inch, y, line[:80])
                y -= 0.2 * inch

            c.save()
            logger.info(f"PDF 文件已保存: {file_path}")

        except ImportError:
            logger.warning("reportlab 未安装，回退为 TXT 格式")
            fallback = file_path.replace(".pdf", ".txt")
            DataExporter.save_analysis_to_txt(analysis_result, fallback, stock_code, analysis_type)

    @staticmethod
    def save_analysis_to_docx(analysis_result: str, file_path: str,
                              stock_code: str = "", analysis_type: str = ""):
        """保存分析结果为 Word 文件"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("财务分析报告", 0)
            doc.add_paragraph(f"股票代码: {stock_code}")
            doc.add_paragraph(f"分析类型: {analysis_type}")
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_heading("分析结果", level=1)

            for line in analysis_result.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

            doc.save(file_path)
            logger.info(f"DOCX 文件已保存: {file_path}")

        except ImportError:
            logger.warning("python-docx 未安装，回退为 TXT 格式")
            fallback = file_path.replace(".docx", ".txt")
            DataExporter.save_analysis_to_txt(analysis_result, fallback, stock_code, analysis_type)
